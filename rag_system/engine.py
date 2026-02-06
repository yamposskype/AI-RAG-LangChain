from __future__ import annotations

import io
import json
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from loguru import logger

from langchain.chains import LLMChain
from langchain.memory import ConversationBufferWindowMemory
from langchain.prompts import PromptTemplate
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.retrievers import EnsembleRetriever
from langchain_community.llms import Ollama
from langchain_community.retrievers import BM25Retriever
from langchain_community.vectorstores import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from sentence_transformers import CrossEncoder

from rag_system.clients import BackendApiClient
from rag_system.config import AppSettings
from rag_system.models import RetrievalResult, RetrievalStrategy
from rag_system.services.agentic_orchestrator import AgenticApiOrchestrator


@dataclass
class RAGConfig:
    """Configuration for the advanced RAG engine."""

    api_base_url: str = "https://rag-langchain-ai-system.onrender.com"
    api_token: str = "token"

    embedding_model: str = "sentence-transformers/all-MiniLM-L6-v2"
    llm_model: str = "llama2"
    rerank_model: str = "cross-encoder/ms-marco-MiniLM-L-6-v2"

    chunk_size: int = 1000
    chunk_overlap: int = 200

    top_k: int = 5
    similarity_threshold: float = 0.7
    enable_reranking: bool = True
    enable_hybrid_search: bool = True

    persist_directory: str = "./chroma_db"
    collection_name: str = "rag_documents"

    memory_key: str = "chat_history"
    max_memory_length: int = 10

    local_documents_dir: str = "backend/documents"

    @classmethod
    def from_settings(cls, settings: AppSettings) -> "RAGConfig":
        return cls(
            api_base_url=settings.api_base_url,
            api_token=settings.api_token,
            embedding_model=settings.embedding_model,
            llm_model=settings.llm_model,
            rerank_model=settings.rerank_model,
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            top_k=settings.top_k,
            similarity_threshold=settings.similarity_threshold,
            enable_reranking=settings.enable_reranking,
            enable_hybrid_search=settings.enable_hybrid_search,
            persist_directory=settings.vector_persist_directory,
            collection_name=settings.vector_collection_name,
            memory_key=settings.memory_key,
            max_memory_length=settings.max_memory_length,
        )


class AdvancedRAGEngine:
    """RAG engine with hybrid retrieval and API-chain enrichment."""

    def __init__(
        self,
        config: Optional[RAGConfig] = None,
        backend_client: Optional[BackendApiClient] = None,
    ):
        self.config = config or RAGConfig()
        self.backend_client = backend_client

        logger.info("Initializing Advanced RAG Engine")
        self.embeddings = self._init_embeddings()
        self.llm = self._init_llm()
        self.reranker = self._init_reranker() if self.config.enable_reranking else None

        self.memory = ConversationBufferWindowMemory(
            memory_key=self.config.memory_key,
            k=self.config.max_memory_length,
            return_messages=False,
        )

        self.vector_store: Optional[Chroma] = None
        self.bm25_retriever: Optional[BM25Retriever] = None
        self.ensemble_retriever: Optional[EnsembleRetriever] = None
        self._indexed_chunks: List[Document] = []
        self.api_orchestrator = (
            AgenticApiOrchestrator(self.backend_client) if self.backend_client else None
        )

    def _init_embeddings(self) -> HuggingFaceEmbeddings:
        logger.info("Loading embedding model: {}", self.config.embedding_model)
        return HuggingFaceEmbeddings(
            model_name=self.config.embedding_model,
            model_kwargs={"device": "cpu"},
            encode_kwargs={"normalize_embeddings": True},
        )

    def _init_llm(self) -> Ollama:
        logger.info("Initializing LLM: {}", self.config.llm_model)
        return Ollama(model=self.config.llm_model, temperature=0.7)

    def _init_reranker(self) -> CrossEncoder:
        logger.info("Loading re-ranker model: {}", self.config.rerank_model)
        return CrossEncoder(self.config.rerank_model)

    @property
    def is_ready(self) -> bool:
        return self.vector_store is not None

    def initialize_from_api(self) -> None:
        """Initialize vector indexes from backend documents API.

        Falls back to local documents when the API is unavailable.
        """
        documents = self.load_documents_from_api()
        if not documents:
            logger.warning(
                "API document load failed; trying local documents at {}",
                self.config.local_documents_dir,
            )
            documents = self.load_documents_from_directory(
                self.config.local_documents_dir
            )

        if not documents:
            logger.error("No documents available to initialize RAG engine")
            return

        self.initialize_with_documents(documents)

    def initialize_with_documents(self, documents: List[Document]) -> None:
        chunks = self.chunk_documents(documents)
        self.build_vector_store(chunks)

        if self.config.enable_hybrid_search:
            self.build_bm25_retriever(chunks)
            self.build_ensemble_retriever()

    def load_documents_from_api(self) -> List[Document]:
        if not self.backend_client:
            return []

        try:
            raw_zip = self.backend_client.download_documents_zip()
        except Exception as exc:
            logger.warning("Unable to download documents from API: {}", exc)
            return []

        return self._extract_documents_from_zip(raw_zip)

    def load_documents_from_directory(self, directory: str) -> List[Document]:
        root = Path(directory)
        if not root.exists():
            return []

        documents: List[Document] = []
        for path in sorted(root.glob("*.txt")):
            try:
                content = path.read_text(encoding="utf-8")
            except Exception as exc:
                logger.warning("Skipping unreadable document {}: {}", path, exc)
                continue
            documents.append(
                Document(
                    page_content=content,
                    metadata={"source": str(path.name), "type": "local_document"},
                )
            )

        logger.info("Loaded {} local documents", len(documents))
        return documents

    def _extract_documents_from_zip(self, zip_bytes: bytes) -> List[Document]:
        documents: List[Document] = []
        with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zip_file:
            for file_name in zip_file.namelist():
                if not file_name.endswith(".txt"):
                    continue
                with zip_file.open(file_name) as file:
                    content = file.read().decode("utf-8")
                documents.append(
                    Document(
                        page_content=content,
                        metadata={"source": file_name, "type": "masterclass_document"},
                    )
                )

        logger.info("Loaded {} documents from API ZIP", len(documents))
        return documents

    def chunk_documents(self, documents: List[Document]) -> List[Document]:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.config.chunk_size,
            chunk_overlap=self.config.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""],
        )

        chunks = text_splitter.split_documents(documents)
        for index, chunk in enumerate(chunks):
            chunk.metadata["chunk_id"] = index
            chunk.metadata["chunk_size"] = len(chunk.page_content)

        logger.info("Created {} document chunks", len(chunks))
        return chunks

    def build_vector_store(self, chunks: List[Document]) -> None:
        Path(self.config.persist_directory).mkdir(parents=True, exist_ok=True)
        self.vector_store = Chroma.from_documents(
            documents=chunks,
            embedding=self.embeddings,
            persist_directory=self.config.persist_directory,
            collection_name=self.config.collection_name,
        )
        self._indexed_chunks = list(chunks)

    def build_bm25_retriever(self, chunks: List[Document]) -> None:
        self.bm25_retriever = BM25Retriever.from_documents(chunks)
        self.bm25_retriever.k = self.config.top_k

    def build_ensemble_retriever(self) -> None:
        if not self.vector_store or not self.bm25_retriever:
            self.ensemble_retriever = None
            return

        semantic_retriever = self.vector_store.as_retriever(
            search_kwargs={"k": self.config.top_k}
        )
        self.ensemble_retriever = EnsembleRetriever(
            retrievers=[semantic_retriever, self.bm25_retriever],
            weights=[0.5, 0.5],
        )

    def add_documents(self, documents: List[Document]) -> int:
        if not documents:
            return 0

        chunks = self.chunk_documents(documents)
        if not chunks:
            return 0

        if self.vector_store is None:
            self.build_vector_store(chunks)
        else:
            self.vector_store.add_documents(chunks)
            self._indexed_chunks.extend(chunks)

        if self.config.enable_hybrid_search:
            self.build_bm25_retriever(self._indexed_chunks)
            self.build_ensemble_retriever()

        return len(chunks)

    def ingest_uploaded_file(self, filepath: str) -> int:
        path = Path(filepath)
        text = self._read_document_text(path)
        if not text.strip():
            return 0

        doc = Document(
            page_content=text,
            metadata={"source": path.name, "type": "uploaded_document"},
        )
        return self.add_documents([doc])

    def _read_document_text(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix in {".txt", ".md"}:
            return path.read_text(encoding="utf-8", errors="ignore")

        if suffix == ".pdf":
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)

        if suffix == ".docx":
            import docx

            document = docx.Document(str(path))
            return "\n".join(paragraph.text for paragraph in document.paragraphs)

        raise ValueError(f"Unsupported file type: {path.suffix}")

    def generate_multi_queries(self, query: str, num_queries: int = 3) -> List[str]:
        template = """You generate alternative retrieval queries.

Original query: {query}
Generate {num_queries} semantically distinct alternatives, one per line."""

        prompt = PromptTemplate(
            template=template, input_variables=["query", "num_queries"]
        )

        try:
            chain = LLMChain(llm=self.llm, prompt=prompt)
            output = self._chain_invoke_text(
                chain, {"query": query, "num_queries": num_queries}
            )
            queries = [
                line.strip("- ").strip() for line in output.splitlines() if line.strip()
            ]
            deduped = [query]
            for item in queries:
                if item not in deduped:
                    deduped.append(item)
            return deduped[:num_queries]
        except Exception as exc:
            logger.warning("Multi-query generation failed: {}", exc)
            return [query]

    def decompose_query(self, query: str) -> List[str]:
        template = """Break the complex query into 2-4 independent sub-queries.

Complex query: {query}
Return one sub-query per line."""

        prompt = PromptTemplate(template=template, input_variables=["query"])

        try:
            chain = LLMChain(llm=self.llm, prompt=prompt)
            output = self._chain_invoke_text(chain, {"query": query})
            parts = [
                line.strip("- ").strip() for line in output.splitlines() if line.strip()
            ]
            return parts or [query]
        except Exception as exc:
            logger.warning("Query decomposition failed: {}", exc)
            return [query]

    def retrieve_documents(
        self,
        query: str,
        strategy: RetrievalStrategy = RetrievalStrategy.HYBRID,
    ) -> RetrievalResult:
        if strategy == RetrievalStrategy.SEMANTIC:
            documents = self._semantic_retrieval(query)
        elif strategy == RetrievalStrategy.HYBRID:
            documents = self._hybrid_retrieval(query)
        elif strategy == RetrievalStrategy.MULTI_QUERY:
            documents = self._multi_query_retrieval(query)
        elif strategy == RetrievalStrategy.DECOMPOSED:
            documents = self._decomposed_retrieval(query)
        else:
            documents = self._semantic_retrieval(query)

        if self.config.enable_reranking and documents:
            documents, scores = self._rerank_documents(query, documents)
        else:
            scores = [1.0] * len(documents)

        return RetrievalResult(
            documents=documents,
            scores=scores,
            query=query,
            strategy=strategy,
            metadata={"num_retrieved": len(documents)},
        )

    def _retrieve_with(self, retriever: Any, query: str) -> List[Document]:
        if hasattr(retriever, "invoke"):
            return list(retriever.invoke(query))
        return list(retriever.get_relevant_documents(query))

    def _semantic_retrieval(self, query: str) -> List[Document]:
        if not self.vector_store:
            return []
        retriever = self.vector_store.as_retriever(
            search_kwargs={"k": self.config.top_k}
        )
        return self._retrieve_with(retriever, query)

    def _hybrid_retrieval(self, query: str) -> List[Document]:
        if not self.ensemble_retriever:
            return self._semantic_retrieval(query)
        return self._retrieve_with(self.ensemble_retriever, query)

    def _multi_query_retrieval(self, query: str) -> List[Document]:
        all_docs: List[Document] = []
        seen = set()

        for q in self.generate_multi_queries(query):
            for doc in self._hybrid_retrieval(q):
                key = hash(doc.page_content)
                if key in seen:
                    continue
                seen.add(key)
                all_docs.append(doc)

        return all_docs[: self.config.top_k * 2]

    def _decomposed_retrieval(self, query: str) -> List[Document]:
        all_docs: List[Document] = []
        seen = set()

        for q in self.decompose_query(query):
            for doc in self._hybrid_retrieval(q):
                key = hash(doc.page_content)
                if key in seen:
                    continue
                seen.add(key)
                all_docs.append(doc)

        return all_docs[: self.config.top_k * 2]

    def _rerank_documents(
        self, query: str, documents: List[Document]
    ) -> Tuple[List[Document], List[float]]:
        if not self.reranker or not documents:
            return documents, [1.0] * len(documents)

        pairs = [[query, doc.page_content] for doc in documents]
        scores = self.reranker.predict(pairs)

        ranked = sorted(zip(documents, scores), key=lambda item: item[1], reverse=True)
        docs = [doc for doc, _ in ranked][: self.config.top_k]
        out_scores = [float(score) for _, score in ranked][: self.config.top_k]

        return docs, out_scores

    def extract_entities(self, text: str) -> Dict[str, List[str]]:
        entities: Dict[str, List[str]] = {
            "persons": [],
            "companies": [],
            "sectors": [],
            "urls": [],
        }

        persons = re.findall(r"\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})\b", text)
        person_blacklist = {
            "Peakspan Masterclasses",
            "Peakspan Masterclass",
            "User Query",
            "Document Context",
        }
        entities["persons"] = [
            person for person in sorted(set(persons)) if person not in person_blacklist
        ]

        companies = re.findall(
            r"\b([A-Z][a-zA-Z0-9&]+(?:\s+[A-Z][a-zA-Z0-9&]+){0,2})\b", text
        )
        company_blacklist = {
            "What",
            "Tell",
            "How",
            "PeakSpan",
            "MasterClass",
            "MasterClasses",
            "Document",
            "Source",
            "Query",
            "Response",
        }
        entities["companies"] = [
            item for item in sorted(set(companies)) if item not in company_blacklist
        ]

        entities["urls"] = re.findall(r"https?://[^\s]+", text)

        sector_keywords = [
            "technology",
            "healthcare",
            "finance",
            "saas",
            "software",
            "fintech",
            "security",
            "enterprise",
            "education",
            "consumer",
        ]
        lower = text.lower()
        entities["sectors"] = [
            keyword for keyword in sector_keywords if keyword in lower
        ]

        return entities

    def call_api_chain(
        self,
        query: str,
        entities: Dict[str, List[str]],
        retrieval_context: str = "",
    ) -> Tuple[Dict[str, Any], List[Dict[str, Any]]]:
        if not self.backend_client or not self.api_orchestrator:
            return {}, []

        llm_seed_calls = self._plan_tool_calls_with_llm(
            query=query, entities=entities, retrieval_context=retrieval_context
        )
        return self.api_orchestrator.execute(
            query=query,
            entities=entities,
            retrieval_context=retrieval_context,
            seed_calls=llm_seed_calls,
        )

    def _plan_tool_calls_with_llm(
        self,
        query: str,
        entities: Dict[str, List[str]],
        retrieval_context: str,
    ) -> List[Dict[str, Any]]:
        if not self.backend_client:
            return []

        tools = self.backend_client.tool_names()
        if not tools:
            return []

        prompt = PromptTemplate(
            template=(
                "You are a planning agent for tool calling.\\n"
                "Available tools: {tools}\\n"
                "Query: {query}\\n"
                "Entities: {entities}\\n"
                "Retrieved Context Snippet: {retrieval_context}\\n\\n"
                "Return ONLY a JSON array of tool calls with this schema:\\n"
                '[{"tool": "team_profile", "params": {"name": "..."}, "reason": "..."}]\\n'
                "Rules:\\n"
                "1) Use only available tools.\\n"
                "2) At most 6 calls.\\n"
                "3) Params must match tool signatures.\\n"
                "4) If no call needed, return []"
            ),
            input_variables=["tools", "query", "entities", "retrieval_context"],
        )

        try:
            chain = LLMChain(llm=self.llm, prompt=prompt)
            raw_output = self._chain_invoke_text(
                chain,
                {
                    "tools": ", ".join(tools),
                    "query": query,
                    "entities": json.dumps(entities),
                    "retrieval_context": retrieval_context[:1200],
                },
            )
            return self._parse_tool_calls_json(raw_output)
        except Exception as exc:
            logger.debug("LLM tool planning failed: {}", exc)
            return []

    @staticmethod
    def _parse_tool_calls_json(raw_output: str) -> List[Dict[str, Any]]:
        text = raw_output.strip()
        if not text:
            return []

        start = text.find("[")
        end = text.rfind("]")
        if start != -1 and end != -1 and end > start:
            text = text[start : end + 1]

        try:
            parsed = json.loads(text)
            if not isinstance(parsed, list):
                return []
            cleaned: List[Dict[str, Any]] = []
            for item in parsed:
                if not isinstance(item, dict):
                    continue
                if "tool" not in item:
                    continue
                params = (
                    item.get("params") if isinstance(item.get("params"), dict) else {}
                )
                cleaned.append(
                    {
                        "tool": str(item["tool"]).strip(),
                        "params": params,
                        "reason": str(item.get("reason", "llm-planned")),
                    }
                )
            return cleaned[:6]
        except Exception:
            return []

    def generate_response(
        self,
        query: str,
        retrieval_result: RetrievalResult,
        api_data: Dict[str, Any],
        api_trace: Optional[List[Dict[str, Any]]] = None,
    ) -> str:
        context_parts = []
        for index, (doc, score) in enumerate(
            zip(retrieval_result.documents, retrieval_result.scores), start=1
        ):
            context_parts.append(
                f"[Document {index}] (Relevance: {score:.2f})\n"
                f"Source: {doc.metadata.get('source', 'Unknown')}\n"
                f"{doc.page_content}\n"
            )

        context = "\n---\n".join(context_parts)
        history = self.memory.load_memory_variables({}).get(self.config.memory_key, "")

        api_summary = ""
        if api_data:
            api_summary = "External Data:\n" + "\n".join(
                f"- {key}: {str(value)[:200]}..." for key, value in api_data.items()
            )
        api_trace_summary = ""
        if api_trace:
            api_trace_summary = "API Chain Trace:\n" + "\n".join(
                f"- {item.get('tool')}({item.get('params')}) => {item.get('status')}"
                for item in api_trace[:12]
            )

        template = """You are an intelligent assistant with access to document context and external API data.

Conversation History:
{history}

Document Context:
{context}

{api_summary}
{api_trace_summary}

User Query: {query}

Instructions:
1. Use the provided evidence.
2. Cite documents as [Document X] when relevant.
3. If information is unavailable, explicitly say so.
4. Keep the answer concise and useful.

Response:"""

        prompt = PromptTemplate(
            template=template,
            input_variables=[
                "history",
                "context",
                "api_summary",
                "api_trace_summary",
                "query",
            ],
        )

        chain = LLMChain(llm=self.llm, prompt=prompt)

        try:
            response = self._chain_invoke_text(
                chain,
                {
                    "history": str(history),
                    "context": context,
                    "api_summary": api_summary,
                    "api_trace_summary": api_trace_summary,
                    "query": query,
                },
            )
            self.memory.save_context({"input": query}, {"output": response})
            return response
        except Exception as exc:
            logger.error("Error generating LLM response: {}", exc)
            return "I apologize, but I encountered an error generating a response. Please try again."

    @staticmethod
    def _chain_invoke_text(chain: LLMChain, payload: Dict[str, Any]) -> str:
        output = chain.invoke(payload)
        if isinstance(output, str):
            return output
        if isinstance(output, dict):
            for key in ("text", "output", "response"):
                value = output.get(key)
                if isinstance(value, str):
                    return value
            values = [value for value in output.values() if isinstance(value, str)]
            if values:
                return values[0]
        return str(output)

    def query(
        self,
        query: str,
        strategy: RetrievalStrategy = RetrievalStrategy.HYBRID,
        use_api_chain: bool = True,
    ) -> Dict[str, Any]:
        retrieval_result = self.retrieve_documents(query, strategy)

        api_data: Dict[str, Any] = {}
        api_trace: List[Dict[str, Any]] = []
        if use_api_chain:
            evidence = " ".join(
                doc.page_content for doc in retrieval_result.documents[:2]
            )
            entities = self.extract_entities(f"{query} {evidence}")
            api_data, api_trace = self.call_api_chain(
                query=query,
                entities=entities,
                retrieval_context=evidence,
            )

        response = self.generate_response(
            query, retrieval_result, api_data, api_trace=api_trace
        )

        return {
            "query": query,
            "response": response,
            "strategy": strategy.value,
            "num_documents": len(retrieval_result.documents),
            "sources": [
                {
                    "source": doc.metadata.get("source", "Unknown"),
                    "score": float(score),
                    "preview": f"{doc.page_content[:200]}...",
                }
                for doc, score in zip(
                    retrieval_result.documents, retrieval_result.scores
                )
            ],
            "api_data_keys": list(api_data.keys()),
            "api_chain_trace": api_trace,
        }
